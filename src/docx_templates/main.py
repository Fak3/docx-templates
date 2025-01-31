from pathlib import Path

from docx import Document
from docx.drawing import Drawing
from jinja2 import Environment, BaseLoader, Undefined
from typer import Option, Typer
from typing_extensions import Annotated



app = Typer(name='docx-templates', pretty_exceptions_enable=True, no_args_is_help=True)


class SilentUndefined(Undefined):
    def _fail_with_undefined_error(self, *args, **kwargs):
        return ''


@app.command()
def main(
    template: Annotated[Path, Option("--template",
        help="Source docx template file",
        exists=True, file_okay=True, dir_okay=False, readable=True, resolve_path=True,
    )],
):
    """
    Populate docx template file with variables from json
    """

    document = Document(template)

    json = {
        "applicationName": "My App",
        "AppFindings": {
            "cvss": {
                "score": 4
            }
        },
        "Paconfig": {
            "applicationName": "My App"
        }
    }

    env = Environment(loader=BaseLoader(), undefined=SilentUndefined)

    def render(block):
        if block.text:
            text = block.text.replace('{{{{', '{{').replace('}}}}', '}}')
            block.text = env.from_string(text).render(json)

    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            for child in run._r.xpath('//*'):
                if child.text and '{{completedOn}}' in child.text:
                    print(child, child.text)
                    import ipdb; ipdb.sset_trace()
                    child.text = '123'
            break
            # for cont in run.iter_inner_content():
                # if hasattr(cont, '_drawing'):
                    # print(cont)
                    # if 'completedOn' in cont._drawing.xml:
                    #     print(cont._drawing.xml)
                # else:
                #     print(cont)
        render(paragraph)
        break

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                # if 'completedOn' in cell.text:
                #     print(cell.text)
                render(cell)

    import ipdb; ipdb.sset_trace()

    document.save('out.docx')


if __name__ == '__main__':
    app()
